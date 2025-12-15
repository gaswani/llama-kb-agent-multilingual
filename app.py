# Copyright (c) 2025 Gideon Aswani / Pathways Technologies Ltd
# Licensed under the MIT License. See LICENSE file for details.

import os
import io
import json
import uuid
import tempfile
from datetime import datetime
from typing import List, Tuple, Dict, Any, Optional
import re

import requests

import numpy as np
import docx  # python-docx

from flask import Flask, request, jsonify, render_template, Response
from flask_cors import CORS

from groq import Groq
from sentence_transformers import SentenceTransformer
from sklearn.metrics.pairwise import cosine_similarity
from sklearn.feature_extraction.text import TfidfVectorizer

# Load .env (recommended)
try:
    from dotenv import load_dotenv
    load_dotenv()
except Exception:
    pass


# =========================
# Configuration
# =========================

GROQ_API_KEY = os.getenv("GROQ_API_KEY")
if not GROQ_API_KEY:
    raise RuntimeError(
        "GROQ_API_KEY not found. Create a .env file with GROQ_API_KEY=your_key_here "
        "or set it in your environment."
    )

LLAMA_MODEL = os.getenv("LLAMA_MODEL", "llama-3.3-70b-versatile")

TOP_K = int(os.getenv("TOP_K", "5"))
SEMANTIC_THRESHOLD = float(os.getenv("SIM_THRESHOLD", "0.40"))

# Hybrid search controls
HYBRID_SEARCH = os.getenv("HYBRID_SEARCH", "true").lower() == "true"
HYBRID_ALPHA = float(os.getenv("HYBRID_ALPHA", "0.70"))        # 0..1 (semantic weight)
HYBRID_THRESHOLD = float(os.getenv("HYBRID_THRESHOLD", "0.35"))  # score floor for hybrid


# =========================
# Language & Location support
# =========================
SUPPORTED_LANGUAGES = {"en": "English", "sw": "Kiswahili"}
DEFAULT_LANGUAGE = os.getenv("DEFAULT_LANGUAGE", "en").strip().lower()
if DEFAULT_LANGUAGE not in SUPPORTED_LANGUAGES:
    DEFAULT_LANGUAGE = "en"

# Reverse geocoding (best-effort) for "near me" queries.
# Requires outbound internet access to Nominatim; if unavailable, app degrades gracefully.
NOMINATIM_URL = os.getenv("NOMINATIM_URL", "https://nominatim.openstreetmap.org/reverse")
NOMINATIM_UA = os.getenv("NOMINATIM_USER_AGENT", "llama-kb-agent/1.0 (contact: admin@example.com)")


# KB persistence
KB_ROOT = os.getenv("KB_ROOT", os.path.join(os.path.dirname(__file__), "kb_store"))
os.makedirs(KB_ROOT, exist_ok=True)
CURRENT_KB_FILE = os.path.join(KB_ROOT, "current_kb.json")


# =========================
# Initialize clients/models
# =========================

llama_client = Groq(api_key=GROQ_API_KEY)
embedder = SentenceTransformer("all-MiniLM-L6-v2")

# In-memory active KB
kb_id_active: Optional[str] = None
kb_chunks: List[str] = []
kb_embeddings: Optional[np.ndarray] = None  # (n_chunks, dim)

# Keyword index (built from kb_chunks)
tfidf_vectorizer: Optional[TfidfVectorizer] = None
tfidf_matrix = None  # sparse


# =========================
# Utility: KB parsing & indexing
# =========================

def load_docx_text(file_bytes: bytes) -> str:
    doc = docx.Document(io.BytesIO(file_bytes))

    parts = []

    # 1) Paragraphs
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)

    # 2) Tables (this is the big missing piece)
    for table in doc.tables:
        for row in table.rows:
            cells = []
            for cell in row.cells:
                ct = (cell.text or "").strip()
                if ct:
                    # Normalize whitespace within a cell
                    ct = " ".join(ct.split())
                    cells.append(ct)
            if cells:
                # Use a separator so keyword search works well
                parts.append(" | ".join(cells))

    # Cleanup common junk
    cleaned = []
    for x in parts:
        if x.lower() in {"bottom of form", "top of form"}:
            continue
        cleaned.append(x)

    return "\n\n".join(cleaned)


def chunk_text(text: str, max_chars: int = 700, overlap: int = 120) -> List[str]:
    """
    Better chunking for Word docs:
    - Splits by blank lines first
    - Treats headings as boundaries (e.g., numbered, ALL CAPS, short title lines)
    - Hard-splits very long blocks with overlap to avoid 1-chunk KB
    """
    # Normalize whitespace
    text = (text or "").replace("\r\n", "\n").replace("\r", "\n").strip()
    blocks = [b.strip() for b in re.split(r"\n\s*\n+", text) if b.strip()]

    def is_heading(line: str) -> bool:
        l = line.strip()
        if not l:
            return False
        # Numbered headings like "1.2 Title" or "3 Title"
        if re.match(r"^\d+(\.\d+)*\s+.+", l):
            return True
        # ALL CAPS short-ish lines
        if len(l) <= 80 and l.upper() == l and any(c.isalpha() for c in l):
            return True
        # Title case short-ish lines (common in docs)
        if len(l) <= 80 and sum(ch.isupper() for ch in l) >= 3 and l.endswith(":"):
            return True
        return False

    # First pass: build sections using headings
    sections: List[str] = []
    buffer: List[str] = []

    for b in blocks:
        lines = [ln.strip() for ln in b.split("\n") if ln.strip()]
        if lines and is_heading(lines[0]) and buffer:
            sections.append("\n\n".join(buffer).strip())
            buffer = [b]
        else:
            buffer.append(b)

    if buffer:
        sections.append("\n\n".join(buffer).strip())

    # Second pass: pack sections into chunks under max_chars
    chunks: List[str] = []
    cur = ""

    def flush():
        nonlocal cur
        if cur.strip():
            chunks.append(cur.strip())
        cur = ""

    for s in sections:
        if len(s) <= max_chars:
            if len(cur) + len(s) + 2 <= max_chars:
                cur = (cur + "\n\n" + s).strip() if cur else s
            else:
                flush()
                cur = s
        else:
            # Hard split long section with overlap
            flush()
            start = 0
            while start < len(s):
                end = min(start + max_chars, len(s))
                piece = s[start:end].strip()
                if piece:
                    chunks.append(piece)
                if end >= len(s):
                    break
                start = max(0, end - overlap)

    flush()

    # Final cleanup: drop tiny chunks
    cleaned = [c for c in chunks if len(c) >= 80]
    return cleaned if cleaned else chunks


def embed_texts(texts: List[str]) -> np.ndarray:
    if not texts:
        return np.empty((0, 0))
    return embedder.encode(texts, convert_to_numpy=True, normalize_embeddings=True)


def build_tfidf_index(chunks: List[str]) -> None:
    """Build TF-IDF matrix for keyword/lexical matching."""
    global tfidf_vectorizer, tfidf_matrix
    if not chunks:
        tfidf_vectorizer, tfidf_matrix = None, None
        return
    tfidf_vectorizer = TfidfVectorizer(
        lowercase=True,
        stop_words="english",
        ngram_range=(1, 2),
        max_features=50000,
    )
    tfidf_matrix = tfidf_vectorizer.fit_transform(chunks)


def semantic_search(query: str, top_k: int = TOP_K) -> Tuple[List[str], List[float]]:
    if kb_embeddings is None or kb_embeddings.shape[0] == 0:
        raise ValueError("Knowledge base is empty. Upload or load a KB first.")

    q_emb = embed_texts([query])[0].reshape(1, -1)
    sims = cosine_similarity(q_emb, kb_embeddings)[0]

    top_indices = np.argsort(-sims)[:top_k]
    return [kb_chunks[i] for i in top_indices], [float(sims[i]) for i in top_indices]


def hybrid_search(
    query: str,
    top_k: int = TOP_K,
    alpha: Optional[float] = None,
) -> Tuple[List[str], List[float], Dict[str, Any]]:
    """
    Hybrid search: blend semantic similarity with TF-IDF cosine similarity.

    score = alpha * semantic + (1 - alpha) * keyword
    """
    if alpha is None:
        alpha = HYBRID_ALPHA

    if kb_embeddings is None or kb_embeddings.shape[0] == 0:
        raise ValueError("Knowledge base is empty. Upload or load a KB first.")

    # Semantic scores
    q_emb = embed_texts([query])[0].reshape(1, -1)
    sem = cosine_similarity(q_emb, kb_embeddings)[0]  # shape: (n,)

    # Keyword scores (TF-IDF cosine similarity)
    kw = None
    if tfidf_vectorizer is not None and tfidf_matrix is not None:
        q_t = tfidf_vectorizer.transform([query])
        kw = (q_t @ tfidf_matrix.T).toarray()[0]  # cosine for L2-normalized tfidf
        # NOTE: scikit TFIDF vectors are L2-normalized by default => dot product ~ cosine.
    else:
        kw = np.zeros_like(sem)

    combined = alpha * sem + (1.0 - alpha) * kw
    top_indices = np.argsort(-combined)[:top_k]

    top_chunks = [kb_chunks[i] for i in top_indices]
    top_scores = [float(combined[i]) for i in top_indices]

    debug = {
        "alpha": alpha,
        "top": [
            {
                "idx": int(i),
                "combined": float(combined[i]),
                "semantic": float(sem[i]),
                "keyword": float(kw[i]),
            }
            for i in top_indices
        ],
    }
    return top_chunks, top_scores, debug


# =========================
# LLM prompting (strict KB grounding)
# =========================

def build_system_prompt(language: str = DEFAULT_LANGUAGE) -> str:
    not_found = "I'm sorry, I couldn't find that in the knowledge base."
    if language == "sw":
        not_found = "Samahani, sijaweza kupata hilo kwenye knowledge base (hati) iliyopakiwa."

    return (
        "You are a helpful, professional AI assistant that answers questions about a single document.\nIMPORTANT: Respond in the user's selected language.\nIf language is 'sw', respond in Kiswahili. If language is 'en', respond in English.\n"
        "You MUST only use the provided DOCUMENT EXCERPTS as your source of truth.\n"
        "- If the answer cannot be derived directly and clearly from the excerpts, you MUST say:\n"
        "  \"I'm sorry, I couldn't find that in the knowledge base.\"\n"
        "- Do NOT invent new facts or rely on outside knowledge.\n"
        "- If the information is ambiguous or incomplete, explain that clearly.\n"
        "- Be concise, well-structured, and easy to read.\n"
).format(NOT_FOUND_MSG=not_found)


def build_user_prompt(question: str, context_chunks: List[str], language: str = DEFAULT_LANGUAGE, location_hint: str = "") -> str:
    context_text = "\n\n---\n\n".join(context_chunks)
    return (
        f"Question:\n{question}\n\n"
        + (f"User location hint (if relevant): {location_hint}\n\n" if location_hint else "")
        + f"DOCUMENT EXCERPTS (knowledge base):\n{context_text}\n\n"
        "Answer the question using ONLY the information in these excerpts."
    )


# =========================
# Swahili -> English translation (for retrieval only)
# =========================
# Why: Your knowledge base is primarily in English (tables/columns/status notes).
# Swahili questions may not retrieve the right chunks using hybrid search.
# We therefore translate the user's Swahili query to English ONLY for retrieval,
# but we still answer in Swahili (language='sw') and still ground answers in KB excerpts.
TRANSLATION_MODEL = os.getenv("TRANSLATION_MODEL", "llama-3.1-8b-instant")

def translate_sw_to_en(text: str) -> str:
    """Translate Swahili text to English for retrieval.

    If translation fails for any reason, we fall back to the original text so the app keeps working.
    """
    t = (text or "").strip()
    if not t:
        return ""

    try:
        resp = llama_client.chat.completions.create(
            model=TRANSLATION_MODEL,
            temperature=0.0,
            max_tokens=256,
            messages=[
                {
                    "role": "system",
                    "content": (
                        "You are a translation engine. Translate Swahili to English. "
                        "Return ONLY the English translation with no quotes, no explanations."
                    ),
                },
                {"role": "user", "content": t},
            ],
        )
        out = resp.choices[0].message.content.strip()
        return out or t
    except Exception:
        return t

def call_llama(question: str, context_chunks: List[str], language: str = DEFAULT_LANGUAGE, location_hint: str = "") -> str:
    completion = llama_client.chat.completions.create(
        model=LLAMA_MODEL,
        messages=[
            {"role": "system", "content": build_system_prompt()},
            {"role": "user", "content": build_user_prompt(question, context_chunks)},
        ],
        temperature=0.1,
        max_tokens=700,
    )
    return completion.choices[0].message.content.strip()


# =========================
# Audio transcription (Groq-hosted Whisper)
# =========================

def transcribe_audio_to_text(file_bytes: bytes, ext: str = "wav", language: str = DEFAULT_LANGUAGE) -> str:
    tmp_path = None
    try:
        # Create temp file path (Windows-safe: close file before reopening)
        fd, tmp_path = tempfile.mkstemp(suffix=f".{ext}")
        with os.fdopen(fd, "wb") as tmp:
            tmp.write(file_bytes)

        # Now reopen for Groq Whisper
        with open(tmp_path, "rb") as f:
            transcription = llama_client.audio.transcriptions.create(
                file=f,
                model="whisper-large-v3-turbo",
                response_format="json",
                temperature=0.0,
                language=language,
            )

        # Depending on Groq SDK response type:
        return transcription.text if hasattr(transcription, "text") else transcription.get("text", "")

    finally:
        if tmp_path and os.path.exists(tmp_path):
            os.remove(tmp_path)


# =========================
# Small-talk / agent behavior
# =========================

def handle_small_talk(query: str, language: str = DEFAULT_LANGUAGE) -> Optional[str]:
    q = (query or "").lower().strip()
    q_no_punct = "".join(ch for ch in q if ch.isalnum() or ch.isspace())

    # English greetings
    greetings = [
        "hi", "hello", "hey",
        "good morning", "good afternoon", "good evening"
    ]

    # Kiswahili greetings
    sw_greetings = [
        "habari", "habari yako", "hujambo", "jambo", "mambo",
        "vipi", "niaje", "salama",
        "za asubuhi", "za mchana", "za jioni",
        "shikamoo"
    ]

    # ---- Greetings ----
    if language == "sw" and any(
        q_no_punct.startswith(g) or f" {g} " in f" {q_no_punct} "
        for g in sw_greetings
    ):
        return (
            "Habari! 👋 Mimi ni msaidizi wako wa AI wa hati. "
            "Najibu maswali kwa kutumia taarifa zilizomo tu kwenye "
            "Digital Hubs – Implementation Status. "
            "Ninaweza kukusaidiaje?"
        )

    if language != "sw" and any(q_no_punct.startswith(g) for g in greetings):
        return (
            "Hello! 👋 I'm your AI document assistant. "
            "I answer questions strictly using the Digital Hubs – Implementation Status "
            "knowledge base. What would you like to know?"
        )

    # ---- Identity / capability ----
    if any(p in q_no_punct for p in [
        "who are you", "what can you do", "what do you do",
        "your role", "your knowledge base"
    ]):
        if language == "sw":
            return (
                "👋 Mimi ni msaidizi wa AI anayejibu maswali kwa kutumia "
                "taarifa zilizomo tu kwenye knowledge base iliyopakiwa. "
                "Uliza kuhusu maeneo, vituo (hubs), hali ya utekelezaji, au taarifa nyingine yoyote."
            )
        return (
            "👋 I'm your AI assistant that answers only from the uploaded knowledge base. "
            "My current knowledge base is the Digital Hubs – Implementation Status. "
            "Ask me about regions, hubs, status, or any other details in the document 😊."
        )

    # ---- Thanks ----
    if "thank" in q_no_punct or "asante" in q_no_punct:
        return (
            "Karibu! 😊 Uliza swali lingine lolote kuhusu Digital Hubs."
            if language == "sw"
            else "You're welcome! 😊 Ask me anything else about Digital Hubs – Implementation Status."
        )

    # ---- Goodbye ----
    if any(p in q_no_punct for p in ["bye", "goodbye", "see you", "see ya", "kwaheri", "tutaonana"]):
        return (
            "Kwaheri! 👋 Rudi wakati wowote ukihitaji msaada."
            if language == "sw"
            else "Goodbye! 👋 Come back anytime you need help with the document."
        )

    return None



# =========================
# KB persistence helpers
# =========================

def kb_dir(kb_id: str) -> str:
    return os.path.join(KB_ROOT, kb_id)


def list_kbs() -> List[str]:
    items = []
    for name in os.listdir(KB_ROOT):
        p = os.path.join(KB_ROOT, name)
        if os.path.isdir(p) and os.path.exists(os.path.join(p, "meta.json")):
            items.append(name)

    def sort_key(k: str) -> str:
        try:
            with open(os.path.join(kb_dir(k), "meta.json"), "r", encoding="utf-8") as f:
                return (json.load(f) or {}).get("created_at", "")
        except Exception:
            return ""
    return sorted(items, key=sort_key, reverse=True)


def save_current_kb_id(kb_id: str) -> None:
    with open(CURRENT_KB_FILE, "w", encoding="utf-8") as f:
        json.dump({"kb_id": kb_id}, f)


def load_current_kb_id() -> Optional[str]:
    if not os.path.exists(CURRENT_KB_FILE):
        return None
    try:
        with open(CURRENT_KB_FILE, "r", encoding="utf-8") as f:
            return (json.load(f) or {}).get("kb_id")
    except Exception:
        return None


def persist_kb(kb_id: str, source_filename: str, source_bytes: bytes, chunks: List[str], embeddings: np.ndarray) -> None:
    d = kb_dir(kb_id)
    os.makedirs(d, exist_ok=True)

    with open(os.path.join(d, "source.docx"), "wb") as f:
        f.write(source_bytes)

    with open(os.path.join(d, "chunks.json"), "w", encoding="utf-8") as f:
        json.dump(chunks, f, ensure_ascii=False, indent=2)

    np.save(os.path.join(d, "embeddings.npy"), embeddings)

    meta = {
        "kb_id": kb_id,
        "source_filename": source_filename,
        "created_at": datetime.utcnow().isoformat() + "Z",
        "num_chunks": len(chunks),
        "embedding_dim": int(embeddings.shape[1]) if embeddings.ndim == 2 and embeddings.shape[0] else 0,
        "llama_model": LLAMA_MODEL,
        "top_k": TOP_K,
        "semantic_threshold": SEMANTIC_THRESHOLD,
        "hybrid_search": HYBRID_SEARCH,
        "hybrid_alpha": HYBRID_ALPHA,
        "hybrid_threshold": HYBRID_THRESHOLD,
    }
    with open(os.path.join(d, "meta.json"), "w", encoding="utf-8") as f:
        json.dump(meta, f, ensure_ascii=False, indent=2)


def load_kb_from_disk(kb_id: str) -> None:
    global kb_id_active, kb_chunks, kb_embeddings
    d = kb_dir(kb_id)
    chunks_path = os.path.join(d, "chunks.json")
    emb_path = os.path.join(d, "embeddings.npy")
    meta_path = os.path.join(d, "meta.json")

    if not (os.path.exists(chunks_path) and os.path.exists(emb_path) and os.path.exists(meta_path)):
        raise ValueError(f"KB '{kb_id}' is missing required files on disk.")

    with open(chunks_path, "r", encoding="utf-8") as f:
        kb_chunks = json.load(f)

    kb_embeddings = np.load(emb_path)
    kb_id_active = kb_id
    save_current_kb_id(kb_id)

    # Build keyword index for hybrid search
    build_tfidf_index(kb_chunks)


def clear_kb_memory() -> None:
    global kb_id_active, kb_chunks, kb_embeddings, tfidf_vectorizer, tfidf_matrix
    kb_id_active = None
    kb_chunks = []
    kb_embeddings = None
    tfidf_vectorizer, tfidf_matrix = None, None


def auto_load_kb_on_startup() -> None:
    candidate = load_current_kb_id()
    if candidate and os.path.isdir(kb_dir(candidate)):
        try:
            load_kb_from_disk(candidate)
            return
        except Exception:
            pass

    versions = list_kbs()
    if versions:
        load_kb_from_disk(versions[0])


# =========================
# Flask app
# =========================

app = Flask(__name__, template_folder="templates", static_folder="static")
CORS(app)

from typing import Any, Dict

def location_hint_from_geo(geo: Dict[str, Any]) -> str:
    """
    Convert a reverse-geocode payload into a short text hint used by retrieval + prompting.
    Expected geo shape (best effort):
      { ok: True, county: "...", state: "...", city: "...", ... }
    """
    if not geo or not isinstance(geo, dict):
        return ""

    # If UI didn’t reverse geocode, you might only have lat/lon; still return something.
    if geo.get("ok") is not True:
        return ""

    parts = []
    if geo.get("county"):
        parts.append(f"County: {geo['county']}")
    if geo.get("state"):
        parts.append(f"Region/State: {geo['state']}")
    if geo.get("city"):
        parts.append(f"City/Town: {geo['city']}")
    if geo.get("display_name") and not parts:
        parts.append(str(geo["display_name"]))

    return "; ".join(parts)

@app.route("/", methods=["GET"])
def ui():
    """
    Serve the Agent UI from templates/index.html (recommended).
    If templates are missing, show a helpful message.
    """
    try:
        return render_template("index.html", active_kb=kb_id_active)
    except Exception:
        return (
            "<h1>Agent is running ✅</h1>"
            "<p>UI template not found. Ensure you have <code>templates/index.html</code> "
            "and <code>static/styles.css</code> in your project.</p>"
            "<p>API endpoints: <code>/upload_kb</code>, <code>/chat_text</code>, <code>/chat</code>, "
            "<code>/kb/list</code>, <code>/kb/load</code>, <code>/kb/reload</code>, <code>/docs</code>.</p>"
        )


@app.route("/health", methods=["GET"])
def health():
    return jsonify({"status": "ok", "active_kb": kb_id_active, "hybrid_search": HYBRID_SEARCH})


# -------- KB management endpoints --------

@app.route("/kb/list", methods=["GET"])
def kb_list():
    return jsonify({"available_kbs": list_kbs(), "current_kb": kb_id_active})


@app.route("/kb/load", methods=["POST"])
def kb_load():
    data = request.get_json(silent=True) or {}
    kb_id = (data.get("kb_id") or "").strip()
    if not kb_id:
        return jsonify({"error": "kb_id is required"}), 400
    try:
        load_kb_from_disk(kb_id)
    except Exception as e:
        return jsonify({"error": str(e)}), 400
    return jsonify({"message": "KB loaded", "current_kb": kb_id_active})


@app.route("/kb/clear", methods=["POST"])
def kb_clear():
    clear_kb_memory()
    return jsonify({"message": "KB cleared from memory", "current_kb": kb_id_active})


@app.route("/kb/reload", methods=["POST"])
def kb_reload():
    if not kb_id_active:
        try:
            auto_load_kb_on_startup()
        except Exception as e:
            return jsonify({"error": str(e)}), 400
        return jsonify({"message": "KB loaded", "current_kb": kb_id_active})

    try:
        load_kb_from_disk(kb_id_active)
    except Exception as e:
        return jsonify({"error": str(e)}), 400
    return jsonify({"message": "KB reloaded", "current_kb": kb_id_active})


@app.route("/upload_kb", methods=["POST"])
def upload_kb():
    if "file" not in request.files:
        return jsonify({"error": "No file part"}), 400

    file = request.files["file"]
    if file.filename == "":
        return jsonify({"error": "No selected file"}), 400

    if not file.filename.lower().endswith(".docx"):
        return jsonify({"error": "Only .docx files are supported"}), 400

    file_bytes = file.read()
    full_text = load_docx_text(file_bytes)
    chunks = chunk_text(full_text, max_chars=500)
    if not chunks:
        return jsonify({"error": "No text found in uploaded document"}), 400

    embeddings = embed_texts(chunks)

    ts = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
    kb_id = f"kb_{ts}_{uuid.uuid4().hex[:8]}"

    persist_kb(kb_id, file.filename, file_bytes, chunks, embeddings)
    load_kb_from_disk(kb_id)

    return jsonify({"message": "Knowledge base uploaded, versioned, and activated.", "kb_id": kb_id_active, "num_chunks": len(kb_chunks)})


# -------- Chat endpoints --------

def retrieve_context(query: str) -> Tuple[List[str], List[float], Dict[str, Any]]:
    """
    Returns: (chunks, scores, debug)
    """
    if HYBRID_SEARCH:
        chunks, scores, dbg = hybrid_search(query, top_k=TOP_K, alpha=HYBRID_ALPHA)
        dbg["mode"] = "hybrid"
        return chunks, scores, dbg
    else:
        chunks, scores = semantic_search(query, top_k=TOP_K)
        return chunks, scores, {"mode": "semantic"}


@app.route("/chat_text", methods=["POST"])
def chat_text():
    data = request.get_json(silent=True) or {}
    message = (data.get("message") or "").strip()
    language = (data.get('language') or DEFAULT_LANGUAGE).strip().lower()
    # Normalize common UI labels
    if language in ("kiswahili", "swahili"):
        language = "sw"
    if language not in SUPPORTED_LANGUAGES:
        language = DEFAULT_LANGUAGE
    geo = data.get('geo') or {}
    location_hint = location_hint_from_geo(geo) if isinstance(geo, dict) else ""
    if not message:
        return jsonify({"error": "Message cannot be empty."}), 400

    small = handle_small_talk(message, language=language)
    if small is not None:
        return jsonify({"answer": small, "query": message, "used_audio": False, "active_kb": kb_id_active, "top_scores": [], "top_snippets": [], "debug": {"mode": "smalltalk"}})

    try:
        # If user selected Kiswahili, translate the query to English for retrieval ONLY.
        retrieval_query = translate_sw_to_en(message) if language == "sw" else message
        context_chunks, scores, dbg = retrieve_context(retrieval_query)
    except ValueError as e:
        return jsonify({"error": str(e)}), 400

    threshold = HYBRID_THRESHOLD if HYBRID_SEARCH else SEMANTIC_THRESHOLD
    if not scores or scores[0] < threshold:
        return jsonify({
            "answer": "I'm sorry, I couldn't find that in the knowledge base.",
            "query": message,
            "used_audio": False,
            "active_kb": kb_id_active,
            "top_scores": scores,
            "top_snippets": context_chunks,
            "debug": dbg,
        })

    answer = call_llama(message, context_chunks, language=language, location_hint=location_hint)
    return jsonify({"answer": answer, "query": message, "used_audio": False, "active_kb": kb_id_active, "top_scores": scores, "top_snippets": context_chunks, "debug": dbg})


@app.route("/chat", methods=["POST"])
def chat():
    used_audio = False
    query = ""

    audio_file = request.files.get("audio")
    language = (request.form.get('language') or DEFAULT_LANGUAGE).strip().lower()
    # Normalize common UI labels
    if language in ("kiswahili", "swahili"):
        language = "sw"
    if language not in SUPPORTED_LANGUAGES:
        language = DEFAULT_LANGUAGE
    geo_raw = request.form.get('geo')
    geo = {}
    if geo_raw:
        try:
            geo = json.loads(geo_raw)
        except Exception:
            geo = {}
    location_hint = location_hint_from_geo(geo) if isinstance(geo, dict) else ""
    if audio_file:
        file_bytes = audio_file.read()
        ext = audio_file.filename.split(".")[-1].lower() if "." in audio_file.filename else "m4a"
        query = transcribe_audio_to_text(file_bytes, ext=ext, language=language)
        used_audio = True
    else:
        query = request.form.get("message", "").strip()

    if not query:
        return jsonify({"error": "Please provide either text or audio."}), 400

    small = handle_small_talk(query, language=language)
    if small is not None:
        return jsonify({"answer": small, "query": query, "used_audio": used_audio, "active_kb": kb_id_active, "top_scores": [], "top_snippets": [], "debug": {"mode": "smalltalk"}})

    try:
        # If user selected Kiswahili, translate the query to English for retrieval ONLY.
        retrieval_query = translate_sw_to_en(query) if language == "sw" else query
        context_chunks, scores, dbg = retrieve_context(retrieval_query)
    except ValueError as e:
        return jsonify({"error": str(e)}), 400

    threshold = HYBRID_THRESHOLD if HYBRID_SEARCH else SEMANTIC_THRESHOLD
    if not scores or scores[0] < threshold:
        return jsonify({
            "answer": "I'm sorry, I couldn't find that in the knowledge base.",
            "query": query,
            "used_audio": used_audio,
            "active_kb": kb_id_active,
            "top_scores": scores,
            "top_snippets": context_chunks,
            "debug": dbg,
        })

    answer = call_llama(query, context_chunks, language=language, location_hint=location_hint)
    return jsonify({"answer": answer, "query": query, "used_audio": used_audio, "active_kb": kb_id_active, "top_scores": scores, "top_snippets": context_chunks, "debug": dbg})


# =========================
# OpenAPI / Swagger UI
# =========================

OPENAPI_SPEC: Dict[str, Any] = {
    "openapi": "3.0.3",
    "info": {
        "title": "Llama KB Agent API",
        "version": "1.0.0",
        "description": "Flask-based AI agent that answers strictly from an uploaded Word (.docx) knowledge base.",
    },
    "servers": [{"url": "http://127.0.0.1:8000", "description": "Local dev"}],
    "paths": {
        "/": {"get": {"summary": "Web UI", "responses": {"200": {"description": "HTML page"}}}},
        "/health": {"get": {"summary": "Health check", "responses": {"200": {"description": "OK"}}}},
        "/upload_kb": {"post": {"summary": "Upload Knowledge Base (.docx)"}},
        "/kb/list": {"get": {"summary": "List KB versions"}},
        "/kb/load": {"post": {"summary": "Load a KB version"}},
        "/kb/clear": {"post": {"summary": "Clear KB from memory"}},
        "/kb/reload": {"post": {"summary": "Reload active KB from disk"}},
        "/chat_text": {"post": {"summary": "Chat (text JSON)"}},
        "/chat": {"post": {"summary": "Chat (multipart: audio or text form)"}},
    },
}



@app.route("/geo/reverse", methods=["POST"])
def geo_reverse():
    data = request.get_json(silent=True) or {}
    try:
        lat = float(data.get("lat"))
        lon = float(data.get("lon"))
    except Exception:
        return jsonify({"error": "lat and lon are required (numbers)."}), 400

    geo = reverse_geocode(lat, lon)
    return jsonify(geo), 200

@app.route("/openapi.json", methods=["GET"])
def openapi_json():
    return jsonify(OPENAPI_SPEC)


@app.route("/docs", methods=["GET"])
def swagger_ui():
    html = """
<!doctype html>
<html>
<head>
  <meta charset="utf-8"/>
  <meta name="viewport" content="width=device-width, initial-scale=1"/>
  <title>API Docs - Llama KB Agent</title>
  <link rel="stylesheet" href="https://unpkg.com/swagger-ui-dist@5/swagger-ui.css">
  <style>body { margin:0; }</style>
</head>
<body>
  <div id="swagger-ui"></div>
  <script src="https://unpkg.com/swagger-ui-dist@5/swagger-ui-bundle.js"></script>
  <script>
    window.onload = () => {
      SwaggerUIBundle({
        url: "/openapi.json",
        dom_id: "#swagger-ui",
        deepLinking: true,
      });
    };
  </script>
</body>
</html>
"""
    return Response(html, mimetype="text/html")


# Auto-load the most recent KB (if any) when the server starts
try:
    auto_load_kb_on_startup()
except Exception:
    pass


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=8000, debug=True)
